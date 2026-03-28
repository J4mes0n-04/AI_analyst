"""Generate specification documents (DOCX / PDF) from project data."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.project_store import ProjectData


def _rp_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def build_docx(data: ProjectData, path: str | Path) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font_name = "Calibri"
    style.font_size = Pt(11)

    doc.add_heading("Техническое задание (черновик)", level=0)
    doc.add_paragraph(
        "Структура ориентирована на рекомендации К. Вигерса: контекст, "
        "бизнес- и пользовательские требования, функциональные требования, трассировка."
    )

    doc.add_heading("1. Контекст и цели", level=1)
    doc.add_paragraph(
        "Кратко опишите продукт, заинтересованные стороны и измеримые цели. "
        "Ниже — исходные бизнес-требования из проекта."
    )
    doc.add_heading("1.1 Исходные бизнес-требования", level=2)
    for para in (data.business_text or "(не заполнено)").split("\n"):
        doc.add_paragraph(para.strip() or " ")

    deep = data.analysis.get("deep_analysis") or {}
    if deep:
        doc.add_heading("1.2 Предпроектный анализ (черновик)", level=2)
        doc.add_paragraph(
            "Сгенерировано эвристически: дополните фактами из исследований и интервью."
        )
        for title, key in [
            ("Проблема и контекст", "problem_and_context"),
            ("Потребности пользователей", "user_needs"),
            ("Рынок", "market_overview"),
            ("Конкуренты", "competitor_landscape"),
            ("Позиционирование", "differentiation_and_positioning"),
            ("Ограничения и комплаенс", "constraints_and_compliance"),
            ("Риски", "risks_and_assumptions"),
            ("Пакеты работ", "recommended_work_packages"),
            ("Открытые вопросы", "open_questions"),
            ("Метрики", "success_metrics"),
        ]:
            block = deep.get(key)
            if not block:
                continue
            doc.add_heading(title, level=3)
            if isinstance(block, list):
                for line in block:
                    if isinstance(line, str):
                        doc.add_paragraph(line, style="List Bullet")
                    else:
                        doc.add_paragraph(str(line), style="List Bullet")
        sh = deep.get("stakeholders")
        if sh:
            doc.add_heading("Стейкхолдеры", level=3)
            for r in sh:
                if isinstance(r, dict):
                    doc.add_paragraph(
                        f"{r.get('role', '')}: интерес — {r.get('interest', '')}; "
                        f"влияние — {r.get('influence', '')}",
                        style="List Bullet",
                    )

    doc.add_heading("2. Пользовательские требования и варианты использования", level=1)
    ucs = data.analysis.get("use_cases") or []
    if not ucs:
        doc.add_paragraph("Варианты использования: сгенерируйте их в модуле анализа.")
    for uc in ucs:
        doc.add_heading(f"{uc.get('id', 'UC')}: {uc.get('name', '')}", level=2)
        doc.add_paragraph(f"Актор: {uc.get('primary_actor', '')}")
        doc.add_paragraph(f"Цель: {uc.get('goal', '')}")
        doc.add_paragraph("Предусловия:")
        for pre in uc.get("preconditions") or []:
            doc.add_paragraph(pre, style="List Bullet")
        doc.add_paragraph("Основной сценарий:")
        for step in uc.get("main_success") or []:
            n = step.get("order", "")
            doc.add_paragraph(
                f"{n}. {step.get('actor_action', '')} → {step.get('system_response', '')}",
                style="List Number",
            )
        doc.add_paragraph("Расширения:")
        for ex in uc.get("extensions") or []:
            doc.add_paragraph(ex, style="List Bullet")

    doc.add_heading("3. Матрица трассировки (сводка)", level=1)
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "ID бизнес-требования"
    hdr[1].text = "Описание"
    hdr[2].text = "User Req"
    hdr[3].text = "Use cases"
    hdr[4].text = "Статус"
    for row in data.matrix_rows:
        cells = table.add_row().cells
        cells[0].text = row.get("business_req_id", "")
        cells[1].text = (row.get("business_text", "") or "")[:500]
        cells[2].text = ", ".join(row.get("user_req_ids") or [])
        cells[3].text = ", ".join(row.get("use_case_ids") or [])
        cells[4].text = row.get("status", "")

    doc.add_heading("4. Приоритизация", level=1)
    for it in data.priority_items[:50]:
        doc.add_paragraph(
            f"[{it.get('band', '')}] {it.get('id', '')}: {it.get('text', '')[:300]} — {it.get('rationale', '')}"
        )

    doc.save(str(path))
    return path


def build_pdf(data: ProjectData, path: str | Path) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, title="Техническое задание")
    styles = getSampleStyleSheet()
    story: list[Any] = []
    story.append(Paragraph("Техническое задание (черновик)", styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(
        Paragraph(
            "Документ сформирован по структуре, согласованной с подходом Вигерса.",
            styles["BodyText"],
        )
    )
    story.append(Spacer(1, 12))
    story.append(Paragraph("1. Бизнес-требования", styles["Heading2"]))
    story.append(Paragraph(_rp_escape((data.business_text or "—")[:8000]), styles["BodyText"]))
    story.append(Spacer(1, 12))
    deep = data.analysis.get("deep_analysis") or {}
    if deep:
        story.append(Paragraph("1.2 Предпроектный анализ (сжато)", styles["Heading2"]))
        for key in (
            "problem_and_context",
            "user_needs",
            "market_overview",
            "competitor_landscape",
            "open_questions",
        ):
            block = deep.get(key)
            if not isinstance(block, list):
                continue
            for line in block[:8]:
                if isinstance(line, str):
                    story.append(Paragraph(_rp_escape(f"• {line[:500]}"), styles["BodyText"]))
        story.append(Spacer(1, 12))
    story.append(Paragraph("2. Варианты использования (кратко)", styles["Heading2"]))
    for uc in (data.analysis.get("use_cases") or [])[:15]:
        line = f"{uc.get('id')}: {uc.get('name', '')} — {uc.get('goal', '')[:200]}"
        story.append(Paragraph(_rp_escape(line), styles["BodyText"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("3. Матрица трассировки", styles["Heading2"]))
    mat = [["BR-ID", "Описание", "UC", "Статус"]]
    for row in data.matrix_rows[:30]:
        mat.append(
            [
                row.get("business_req_id", ""),
                (row.get("business_text", "") or "")[:80],
                ", ".join(row.get("use_case_ids") or []),
                row.get("status", ""),
            ]
        )
    t = Table(mat, colWidths=[60, 220, 80, 60])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e2e8f0")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.append(t)
    doc.build(story)
    path.write_bytes(buf.getvalue())
    return path
